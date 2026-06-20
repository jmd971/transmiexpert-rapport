"""
═══════════════════════════════════════════════════════════════════════════════
SAGETRIM — GÉNÉRATEUR DE RAPPORT WORD (.docx) ÉDITABLE
Équivalent éditable du template PDF — Luc SILVESTRE
═══════════════════════════════════════════════════════════════════════════════

USAGE :
    from sagetrim_docx import RapportExpertise
    RapportExpertise(params).generer("/tmp/rapport.docx")

Reproduit la charte SAGETRIM (navy/or) en document Word modifiable.
Même interface publique et mêmes paramètres que le template PDF.
"""

import io
import re
import base64
import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── PALETTE SAGETRIM (hex sans #) ───────────────────────────────────────────
NAVY="1A2B45"; GOLD="C8A96E"; LIGHT="F7F5F0"; LIGHTD="EDE9E1"; BORDER="D0C8B8"
TEXT="1A1712"; MUTED="6B6355"; ALERT="8B2020"; ORANGE="C8682A"; GREEN="1A4A2E"
WHITE="FFFFFF"; BLUEBG="EEF3FA"; GREENBG="EAF0E8"; WARNBG="FDF6E3"; ORANGEBG="FFF3E6"

def RGB(h): return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

FONT_BODY = "Calibri"
FONT_HEAD = "Cambria"

# ─── CONSTANTES CABINET ──────────────────────────────────────────────────────
CABINET = {
    "nom"      : "SAGETRIM",
    "sous_nom" : "Cabinet Martial — Syndic de copropriétés — Administrateur de Biens",
    "expert"   : "Luc SILVESTRE",
    "titre"    : "Expert en évaluation immobilière",
    "adresse"  : "37 Espace Rocade — Grand-Camp — 97139 Les Abymes (Guadeloupe)",
    "tel"      : "0590 82 78 76",
    "rc"       : "RC N° 77B22",
    "siret"    : "31004861600026",
    "ape"      : "6831Z",
    "capital"  : "13 720,41 €",
    "rcp"      : "Sévères Assurance — Garantie financière : 500 000 €",
    "norme"    : ("Charte de l'Expertise en Évaluation Immobilière (5ᵉ éd. 2019) "
                  "et normes TEGoVA (EVS 2022)"),
    "refs_legales": (
        "Code civil art. 1075 (donation-partage) — "
        "Code de l'urbanisme art. R. 111-22 (Surface de Plancher — SDP, "
        "en vigueur depuis le 1ᵉʳ mars 2012 — décret 2011-2054 abrogeant SHOB/SHON) — "
        "CCH art. L. 271-4 (diagnostics obligatoires) — "
        "Base DVF+ open-data (DGFiP / Cerema — data.gouv.fr) — "
        "Géorisques (georisques.gouv.fr)"
    ),
}

CW = Cm(16.8)   # largeur utile (A4 21cm - marges 2.2 + 2.0)


# ═══════════════════════════════════════════════════════════════════════════════
# HELPERS XML BAS NIVEAU
# ═══════════════════════════════════════════════════════════════════════════════
def _shade(props, fill):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    props.append(shd)

def shade_cell(cell, fill):
    _shade(cell._tc.get_or_add_tcPr(), fill)

def shade_para(p, fill):
    _shade(p._p.get_or_add_pPr(), fill)

def _border_el(tag, color, sz, space=1):
    e = OxmlElement(tag)
    e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(sz))
    e.set(qn('w:space'), str(space)); e.set(qn('w:color'), color)
    return e

def para_border(p, edge, color, sz=12, space=2):
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr'); pPr.append(pBdr)
    pBdr.append(_border_el(f'w:{edge}', color, sz, space))

def cell_border(cell, edge, color, sz=12):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders'); tcPr.append(borders)
    borders.append(_border_el(f'w:{edge}', color, sz))

def table_borders(table, color=BORDER, sz=4, inside=True, outside=True):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    edges = []
    if outside: edges += ['top', 'left', 'bottom', 'right']
    if inside:  edges += ['insideH', 'insideV']
    for edge in edges:
        borders.append(_border_el(f'w:{edge}', color, sz))
    tblPr.append(borders)

def set_col_widths(table, widths):
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = w

def cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('bottom', bottom), ('start', left), ('end', right)):
        e = OxmlElement(f'w:{tag}')
        e.set(qn('w:w'), str(val)); e.set(qn('w:type'), 'dxa')
        m.append(e)
    tcPr.append(m)

def add_page_number(paragraph):
    run = paragraph.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    run._r.append(f1); run._r.append(instr); run._r.append(f2)
    run.font.size = Pt(8); run.font.name = FONT_BODY; run.font.color.rgb = RGB(MUTED)


# ═══════════════════════════════════════════════════════════════════════════════
# RUNS / TEXTE RICHE (gère <b>…</b> et les retours ligne \n)
# ═══════════════════════════════════════════════════════════════════════════════
def add_rich(p, text, size=10, color=TEXT, italic=False, font=FONT_BODY, bold_base=False):
    if text is None:
        text = ""
    text = str(text)
    bold = bold_base
    for seg in re.split(r'(</?b>)', text):
        if seg == '<b>':  bold = True;  continue
        if seg == '</b>': bold = False; continue
        if seg == '':     continue
        lines = seg.split('\n')
        for k, line in enumerate(lines):
            if k > 0:
                p.add_run().add_break()
            r = p.add_run(line)
            r.font.size = Pt(size); r.bold = bold; r.italic = italic
            r.font.color.rgb = RGB(color); r.font.name = font


# ═══════════════════════════════════════════════════════════════════════════════
# CLASSE PRINCIPALE
# ═══════════════════════════════════════════════════════════════════════════════
class RapportExpertise:
    """Générateur de rapport d'expertise SAGETRIM au format Word éditable."""

    def __init__(self, params: dict):
        self.p = self._valider(params)
        self.doc = Document()
        self._reserves = []
        self._setup_document()

    # ── VALIDATION ────────────────────────────────────────────────────────────
    def _valider(self, p):
        type_bien = p.get("type_bien", "maison")
        labels = {"maison": "Maison individuelle",
                  "appartement": "Appartement en copropriété",
                  "terrain": "Terrain nu"}
        p["type_bien"] = type_bien
        p["type_bien_label"] = labels.get(type_bien, type_bien)
        p.setdefault("hypothese", "Valeur vénale — pleine propriété — hors fiscalité")
        p.setdefault("objet_mission", "Estimation en valeur vénale")
        p.setdefault("zone_sismique", "V")
        p.setdefault("assainissement", "Collectif")
        p.setdefault("reserves_initiales", "Néant sauf indication contraire")
        if p.get("type_bien") == "maison" and not p.get("vente_anterieure"):
            _prix = float(p.get("origine_prix_acquisition") or 0)
            _ib = float(p.get("origine_indice_icc") or 0)
            _ir = float(p.get("origine_indice_icc_actuel") or 0)
            if _prix > 0 and _ib > 0 and _ir > 0:
                p["vente_anterieure"] = {
                    "prix": _prix, "date": p.get("origine_date_acquisition", ""),
                    "indice_base": _ib, "indice_revision": _ir,
                    "trimestre_base": "trimestre d'acquisition",
                    "trimestre_revision": "dernier trimestre publié",
                }
        return p

    # ── MISE EN PAGE / EN-TÊTE / PIED ─────────────────────────────────────────
    def _setup_document(self):
        sec = self.doc.sections[0]
        sec.page_height = Cm(29.7); sec.page_width = Cm(21)
        sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.0)
        sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)

        style = self.doc.styles['Normal']
        style.font.name = FONT_BODY
        style.font.size = Pt(10)
        style.font.color.rgb = RGB(TEXT)
        style.paragraph_format.space_after = Pt(4)

        # Pas d'en-tête/pied sur la page de garde
        sec.different_first_page_header_footer = True

        p = self.p
        # En-tête (pages 2+)
        hp = sec.header.paragraphs[0]
        hp.text = ""
        add_rich(hp, f"SAGETRIM — {CABINET['expert']}", size=8, bold_base=True, color=NAVY)
        hp.add_run().add_tab()
        ref_line = f"{p.get('ref','')}  •  {(p.get('demandeur_nom','') or '').split()[0]}  •  {p.get('date_rapport','')}"
        add_rich(hp, ref_line, size=8, color=MUTED)
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para_border(hp, 'bottom', GOLD, sz=12, space=2)

        # Pied (pages 2+)
        fp = sec.footer.paragraphs[0]
        fp.text = ""
        add_rich(fp, f"Rapport établi conformément à la {CABINET['norme']}.", size=7, italic=True, color=MUTED)
        fp.add_run().add_tab()
        add_rich(fp, "Page ", size=8, color=MUTED)
        add_page_number(fp)
        para_border(fp, 'top', BORDER, sz=4, space=2)

    # ── RÉSERVES AUTOMATIQUES ─────────────────────────────────────────────────
    def _ajouter_reserve(self, texte, niveau="alerte"):
        self._reserves.append((texte, niveau))

    def _collecter_reserves(self):
        diags = {
            "diag_amiante" : "Diagnostic amiante non fourni (obligatoire si construction avant 1997)",
            "diag_termites": "Diagnostic termites non fourni (obligatoire en zone contaminée)",
            "diag_elec"    : "Diagnostic électricité non fourni (obligatoire si installation > 15 ans)",
            "diag_dpe"     : "DPE non fourni — classe énergétique inconnue",
        }
        for key, msg in diags.items():
            if not self.p.get(key, False):
                self._ajouter_reserve(msg)
        if not self.p.get("sdp"):
            self._ajouter_reserve("Surface de Plancher (SDP) non renseignée — valeur indicative", "info")

    # ── FORMAT ────────────────────────────────────────────────────────────────
    def _fmt(self, n):
        return f"{n:,.0f}".replace(",", " ") + " €"

    # ═════════════════════════════════════════════════════════════════════════
    # COMPOSANTS RÉUTILISABLES
    # ═════════════════════════════════════════════════════════════════════════
    def section(self, titre):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(14); pf.space_after = Pt(8)
        pf.left_indent = Pt(8)
        shade_para(p, NAVY)
        para_border(p, 'bottom', GOLD, sz=18, space=2)
        add_rich(p, titre, size=13, bold_base=True, color=WHITE, font=FONT_HEAD)
        return p

    def h2(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
        add_rich(p, text, size=11, bold_base=True, color=NAVY, font=FONT_HEAD)
        return p

    def h3(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(3)
        add_rich(p, text, size=10, bold_base=True, color=NAVY)
        return p

    def para(self, text, size=10, color=TEXT, italic=False, align='justify', space_after=6):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        a = {'justify': WD_ALIGN_PARAGRAPH.JUSTIFY, 'right': WD_ALIGN_PARAGRAPH.RIGHT,
             'center': WD_ALIGN_PARAGRAPH.CENTER, 'left': WD_ALIGN_PARAGRAPH.LEFT}
        p.alignment = a.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        add_rich(p, text, size=size, color=color, italic=italic)
        return p

    def bullet(self, text):
        p = self.doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        add_rich(p, text, size=10, color=TEXT)
        return p

    def kv(self, rows, c1=Cm(6)):
        table = self.doc.add_table(rows=0, cols=2)
        widths = [c1, CW - c1]
        for i, (k, v) in enumerate(rows):
            cells = table.add_row().cells
            ck = cells[0].paragraphs[0]; ck.paragraph_format.space_after = Pt(0)
            add_rich(ck, k, size=9, bold_base=True, color=MUTED)
            cv = cells[1].paragraphs[0]; cv.paragraph_format.space_after = Pt(0)
            add_rich(cv, v, size=10, color=TEXT)
            bg = LIGHT if i % 2 else WHITE
            for c in cells:
                shade_cell(c, bg); cell_margins(c)
                cell_border(c, 'bottom', BORDER, sz=2)
        set_col_widths(table, widths)
        self._spacer(4)
        return table

    def legal(self, texte):
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        shade_cell(cell, BLUEBG); cell_margins(cell, top=80, bottom=80, left=140)
        cell_border(cell, 'left', NAVY, sz=18)
        cp = cell.paragraphs[0]; cp.paragraph_format.space_after = Pt(0)
        add_rich(cp, texte, size=8, color=TEXT)
        set_col_widths(table, [CW])
        self._spacer(4)
        return table

    def reserve(self, texte, couleur=ALERT, bg=WARNBG):
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        shade_cell(cell, bg); cell_margins(cell, top=80, bottom=80, left=140)
        cell_border(cell, 'left', couleur, sz=18)
        cp = cell.paragraphs[0]; cp.paragraph_format.space_after = Pt(0)
        add_rich(cp, f"Réserve : {texte}", size=9, bold_base=True, color=couleur)
        set_col_widths(table, [CW])
        self._spacer(4)
        return table

    def conclusion_box(self, val_min, val_ret, val_max):
        fmt = self._fmt
        table = self.doc.add_table(rows=4, cols=2)
        c1 = Cm(7.5); widths = [c1, CW - c1]
        # Ligne 0 — bandeau navy
        r0 = table.rows[0].cells
        p0 = r0[0].paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER; p0.paragraph_format.space_after = Pt(0)
        add_rich(p0, "VALEUR VÉNALE RETENUE", size=11, bold_base=True, color=WHITE)
        p1 = r0[1].paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER; p1.paragraph_format.space_after = Pt(0)
        add_rich(p1, fmt(val_ret), size=22, bold_base=True, color=GOLD, font=FONT_HEAD)
        for c in r0:
            shade_cell(c, NAVY); cell_margins(c, top=120, bottom=120)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Lignes fourchette + note
        data = [("Fourchette basse", fmt(val_min)),
                ("Fourchette haute", fmt(val_max)),
                ("Hors fiscalité — hors frais d'acte", "")]
        for i, (lab, val) in enumerate(data, start=1):
            cells = table.rows[i].cells
            pa = cells[0].paragraphs[0]; pa.paragraph_format.space_after = Pt(0)
            add_rich(pa, lab, size=8, color=MUTED, italic=(i == 3))
            pb = cells[1].paragraphs[0]; pb.paragraph_format.space_after = Pt(0)
            add_rich(pb, val, size=8, color=MUTED)
            for c in cells:
                shade_cell(c, LIGHT); cell_margins(c)
        set_col_widths(table, widths)
        table_borders(table, color=GOLD, sz=12, inside=False, outside=True)
        self._spacer(6)
        return table

    # ── Tableaux génériques ────────────────────────────────────────────────────
    def _styled_table(self, header, rows, widths, total_row=False):
        n = len(header)
        table = self.doc.add_table(rows=1, cols=n)
        hcells = table.rows[0].cells
        for j, h in enumerate(header):
            hp = hcells[j].paragraphs[0]; hp.paragraph_format.space_after = Pt(0)
            add_rich(hp, h, size=8, bold_base=True, color=WHITE)
            shade_cell(hcells[j], NAVY); cell_margins(hcells[j])
        nrows = len(rows)
        for ri, row in enumerate(rows):
            cells = table.add_row().cells
            is_total = total_row and ri == nrows - 1
            for j, val in enumerate(row):
                cp = cells[j].paragraphs[0]; cp.paragraph_format.space_after = Pt(0)
                add_rich(cp, val, size=8, color=TEXT, bold_base=is_total)
                if is_total:
                    shade_cell(cells[j], GREENBG)
                else:
                    shade_cell(cells[j], LIGHT if ri % 2 else WHITE)
                cell_margins(cells[j])
                if is_total:
                    cell_border(cells[j], 'top', GOLD, sz=12)
        set_col_widths(table, widths)
        table_borders(table, color=BORDER, sz=3)
        self._spacer(6)
        return table

    def dvf_table(self, refs, pm2_median):
        header = ["Date", "Localisation / Résidence", "Type", "SDP", "Valeur", "€/m²", "Statut"]
        rows = []
        for r in refs:
            retenu = r.get("retenu", True)
            pm2 = r.get("pm2", 0)
            statut = r.get("statut", ("Retenue" if retenu else "Écartée"))
            rows.append([
                r.get("date", ""),
                r.get("localisation", ""),
                r.get("type", ""),
                f"{r.get('surface', 0)} m²",
                f"{r.get('valeur', 0):,.0f}".replace(",", " ") + " €",
                f"{pm2:,.0f}".replace(",", " ") + " €",
                statut,
            ])
        nb = sum(1 for r in refs if r.get('retenu', True))
        rows.append(["<b>MÉDIANE RETENUE</b>", f"{nb} références", "", "", "",
                     f"<b>{pm2_median:,.0f}".replace(",", " ") + " €/m²</b>", "Base de calcul"])
        widths = [Cm(2.2), Cm(4.0), Cm(1.5), Cm(1.8), Cm(2.8), Cm(2.5), Cm(2.0)]
        return self._styled_table(header, rows, widths, total_row=True)

    def risque_table(self, risques):
        header = ["Risque", "Niveau", "Impact valeur", "Source"]
        widths = [Cm(3.5), Cm(2.5), Cm(7.8), Cm(3.0)]
        return self._styled_table(header, risques, widths)

    def synth_table(self, methodes):
        header = ["Méthode", "Valeur", "Pondération", "Contribution"]
        widths = [Cm(7.0), Cm(3.3), Cm(3.0), Cm(3.5)]
        return self._styled_table(header, methodes, widths, total_row=True)

    def _vetuste_table(self, postes):
        header = ["Poste", "% coût", "Âge eff.", "Durée vie", "Vét. poste", "Contrib."]
        rows = []
        total = 0.0
        for p in postes:
            age = float(p.get("age_effectif", 0) or 0)
            dv  = float(p.get("duree_vie", 1) or 1)
            pct = float(p.get("pct_cout", 0) or 0)
            vet = min(age / dv, 1.0) if dv else 0
            contrib = pct * vet
            total += contrib
            rows.append([p.get("poste", ""), f"{pct:.0f} %", f"{age:.0f} ans",
                         f"{dv:.0f} ans", f"{vet*100:.0f} %", f"{contrib:.1f} %"])
        rows.append(["<b>VÉTUSTÉ GLOBALE PONDÉRÉE</b>", "", "", "", "", f"<b>{total:.1f} %</b>"])
        widths = [Cm(6.5), Cm(1.7), Cm(2.0), Cm(2.0), Cm(2.0), Cm(2.6)]
        return self._styled_table(header, rows, widths, total_row=True)

    def _carac_physiques_table(self, postes):
        header = ["Poste", "État constaté", "Observations"]
        rows = [[c.get("poste", ""), c.get("etat", ""), c.get("notes", "")] for c in postes]
        widths = [Cm(4.5), Cm(3.3), Cm(9.0)]
        return self._styled_table(header, rows, widths)

    # ── PHOTOS ─────────────────────────────────────────────────────────────────
    def _embed_photos(self, photos):
        from docx.shared import Cm as _Cm
        img_w = Cm(7.8)
        for i in range(0, len(photos), 2):
            pair = photos[i:i+2]
            table = self.doc.add_table(rows=2, cols=2)
            for j in range(2):
                img_cell = table.rows[0].cells[j]
                cap_cell = table.rows[1].cells[j]
                ip = img_cell.paragraphs[0]; ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ip.paragraph_format.space_after = Pt(0)
                cp = cap_cell.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.paragraph_format.space_after = Pt(0)
                if j < len(pair):
                    photo = pair[j]
                    data = photo.get("data", "")
                    caption = photo.get("caption", "") or photo.get("name", "")
                    try:
                        if isinstance(data, bytes):
                            raw = data
                        elif isinstance(data, str):
                            raw = base64.b64decode(data.split(",", 1)[1] if "," in data else data)
                        else:
                            raise ValueError("type photo non reconnu")
                        ip.add_run().add_picture(io.BytesIO(raw), width=img_w)
                    except Exception as e:
                        add_rich(ip, "[Photo indisponible]", size=8, color=MUTED, italic=True)
                    add_rich(cp, caption, size=8, color=MUTED, italic=True)
            set_col_widths(table, [Cm(8.2), Cm(8.2)])
            self._spacer(6)

    def _spacer(self, h=6):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = Pt(h)
        return p

    # ═════════════════════════════════════════════════════════════════════════
    # PAGE DE GARDE
    # ═════════════════════════════════════════════════════════════════════════
    def _cover(self):
        p = self.p
        # Bandeau navy
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        shade_cell(cell, NAVY); cell_margins(cell, top=200, bottom=200, left=120, right=120)
        set_col_widths(table, [CW])

        def cline(text, size, color, font=FONT_BODY, after=2, bold=True):
            cp = cell.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(after); cp.paragraph_format.space_before = Pt(0)
            add_rich(cp, text, size=size, color=color, font=font, bold_base=bold)
            return cp

        # premier paragraphe de la cellule (déjà présent)
        first = cell.paragraphs[0]
        first.alignment = WD_ALIGN_PARAGRAPH.CENTER
        first.paragraph_format.space_after = Pt(2)
        add_rich(first, CABINET["nom"], size=24, bold_base=True, color=WHITE, font=FONT_HEAD)
        cline(CABINET["sous_nom"], 10, WHITE, bold=False, after=2)
        cline("Guadeloupe", 9, GOLD, after=10, bold=False)
        cline("RAPPORT D'EXPERTISE", 22, WHITE, font=FONT_HEAD, after=2)
        cline("Évaluation en valeur vénale", 12, WHITE, bold=False, after=4)

        self._spacer(14)

        # Bloc bien
        self.para("BIEN EXPERTISÉ", size=10, color=NAVY, align='center', space_after=4)
        bien_titre = (p.get("adresse_bien", p.get("commune", "")) or "").upper()
        bp = self.doc.add_paragraph(); bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bp.paragraph_format.space_after = Pt(2)
        add_rich(bp, bien_titre, size=13, bold_base=True, color=TEXT, font=FONT_HEAD)
        sous = f"{p.get('type_bien_label','Bien immobilier')} — {p.get('commune','')} ({p.get('code_postal','')})"
        self.para(sous, size=10, color=TEXT, align='center', space_after=8)

        # Bloc demandeur (encadré clair)
        dtable = self.doc.add_table(rows=1, cols=1)
        dcell = dtable.cell(0, 0); shade_cell(dcell, LIGHT); cell_margins(dcell, top=100, bottom=100, left=140)
        set_col_widths(dtable, [CW])
        d1 = dcell.paragraphs[0]; d1.paragraph_format.space_after = Pt(2)
        add_rich(d1, "Demandeur : ", size=9, bold_base=True, color=NAVY)
        add_rich(d1, p.get("demandeur_nom", ""), size=9, color=TEXT)
        d2 = dcell.add_paragraph(); d2.paragraph_format.space_after = Pt(0)
        add_rich(d2, f"Mission : {p.get('objet_mission','Estimation en valeur vénale')}", size=9, color=TEXT)
        self._spacer(8)

        # Infos
        self.kv([
            ("Référence dossier", p.get("ref", "")),
            ("Date de visite",    p.get("date_visite", "")),
            ("Date du rapport",   p.get("date_rapport", "")),
            ("Hypothèse",         p.get("hypothese", "Valeur vénale — pleine propriété — hors fiscalité")),
        ])
        self._spacer(10)

        # Pied de page de garde
        rule = self.doc.add_paragraph(); rule.paragraph_format.space_after = Pt(4)
        para_border(rule, 'bottom', BORDER, sz=4, space=1)
        self.para(CABINET["expert"], size=9, color=NAVY, align='left', space_after=1)
        self.para(f"{CABINET['titre']}  •  {CABINET['adresse']}", size=8, color=MUTED, align='left', space_after=1)
        self.para(f"{CABINET['rc']}  •  SIRET {CABINET['siret']}  •  RCP : {CABINET['rcp']}", size=8, color=MUTED, align='left', space_after=1)
        self.para(f"Tél. {CABINET['tel']}", size=8, color=MUTED, align='left', space_after=1)

        self.doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # SECTIONS DU RAPPORT
    # ═════════════════════════════════════════════════════════════════════════
    def _section_I_serment(self):
        p = self.p
        self.section("I  —  PRESTATION DE SERMENT")
        self.para(
            f"Je soussigné, <b>{CABINET['expert']}</b>, {CABINET['titre']}, "
            f"domicilié au {CABINET['adresse']}, avoir pris acte de la mission "
            f"qui m'a été confiée par <b>{p.get('demandeur_nom','')}</b>.")
        self.para("Je déclare accepter la présente mission.")
        self.para("Je m'engage à la remplir fidèlement, consciencieusement et impartialement, "
                  f"conformément aux dispositions de la <b>{CABINET['norme']}</b>.")
        self.para(
            "Je certifie n'avoir aucun lien d'intérêt, direct ou indirect, avec le bien expertisé "
            "ou avec les parties, de nature à compromettre mon indépendance ou mon objectivité. La présente "
            "expertise a été conduite en toute indépendance ; les honoraires y afférents ne sont liés ni au "
            "résultat de l'évaluation ni à la conclusion du rapport.")
        self._spacer(16)
        self.para(CABINET["nom"], size=11, color=NAVY, align='right', space_after=2)
        self.para(CABINET["expert"], size=11, color=NAVY, align='right', space_after=4)

    def _section_II_clause(self):
        self.section("II  —  CLAUSE SPÉCIFIQUE DE CONFIDENTIALITÉ")
        self.para(
            "Il est formellement convenu qu'il est interdit au client d'utiliser le présent "
            "rapport d'expertise de façon partielle en isolant telle ou telle partie de son contenu.")
        self.para(
            "Le présent rapport d'expertise, en tout ou partie, ne pourra être cité ni même "
            "mentionné dans aucun document, aucune circulaire et aucune déclaration destinés à être "
            "publiés, ne pourra l'être d'une quelconque manière sans l'accord écrit de l'expert "
            "quant à la forme et aux circonstances dans lesquelles il pourra paraître.")
        self.legal(
            f"Ce rapport est établi en deux exemplaires originaux. "
            f"Réf. dossier : {self.p.get('ref','')}  •  Expert : {CABINET['expert']}  •  "
            f"RCP : {CABINET['rcp']}")

    def _section_III_mission(self):
        p = self.p
        docs = p.get("documents_transmis", ["Titre de propriété"])
        docs_str = "\n".join(f"• {d}" for d in docs) if docs else "Titre de propriété"
        self.section("III  —  MISSION")
        self.h2("1  —  Identité du demandeur")
        self.kv([
            ("Demandeur", p.get("demandeur_nom", "")),
            ("Qualité",   p.get("demandeur_qualite", "")),
            ("Adresse",   p.get("demandeur_adresse", "")),
        ])
        self.h2("2  —  Objet de la mission")
        self.kv([
            ("Mission",   p.get("objet_mission", "")),
            ("Bien",      f"{p.get('type_bien_label','')} — {p.get('adresse_bien','')}"),
            ("Hypothèse", p.get("hypothese", "")),
            ("Réserves",  p.get("reserves_initiales", "Néant")),
        ])
        self.h2("3  —  Visite et documents transmis")
        self.kv([
            ("Date de visite",    p.get("date_visite", "")),
            ("Présence",          p.get("accompagnateur", p.get("demandeur_nom", ""))),
            ("Documents fournis", docs_str),
        ])
        self.legal(
            f"Référence dossier : {p.get('ref','')}  •  "
            f"Date du rapport : {p.get('date_rapport','')}  •  "
            f"Méthodes : {self._label_methodes()}")
        self.h2("4  —  Base de la valeur et date d'effet")
        self.para(
            "La valeur recherchée est la <b>valeur vénale</b>, définie comme le montant estimé auquel le bien "
            "devrait s'échanger, à la date de l'évaluation, entre un acheteur et un vendeur consentants, dans "
            "le cadre d'une transaction équilibrée conclue dans des conditions normales de concurrence, après "
            "une commercialisation adéquate, où chacune des parties a agi en connaissance de cause, prudemment "
            "et sans contrainte.")
        _dval = p.get("date_visite") or p.get("date_rapport", "")
        self.para(
            f"La valeur est arrêtée à la date de nos constatations sur place, le <b>{_dval}</b>. Elle s'entend "
            "hors droits de mutation, frais d'acte et fiscalité éventuelle, pour la pleine propriété d'un bien "
            "libre de toute occupation, sauf mention contraire au présent rapport.")
        self.h2("5  —  Étendue et limites des investigations")
        self.para(
            "Nos investigations ont consisté en une visite des lieux, l'examen des documents qui nous ont été "
            "communiqués et l'analyse des données de marché disponibles. Elles n'ont pas comporté de sondage "
            "destructif, de vérification de la solidité des structures, ni de relevé métré par géomètre : les "
            "surfaces indiquées sont reprises des documents fournis ou de nos constatations, à titre indicatif.")
        self.para(
            "Nous avons présumé exactes et sincères les informations et pièces transmises par le demandeur, "
            "sans en garantir l'authenticité juridique. Le présent rapport ne se substitue ni aux diagnostics "
            "techniques réglementaires, ni à l'analyse notariale du titre de propriété.")

    def _label_methodes(self):
        t = self.p.get("type_bien", "maison")
        if t == "appartement":
            return "Comparaison directe (DVF)"
        elif t == "terrain":
            return "Méthode terrain"
        else:
            methodes = ["Comparaison directe (DVF)", "Sol & Construction"]
            if self.p.get("vente_anterieure"):
                methodes.append("Actualisation par indice ICC")
            return " + ".join(methodes)

    def _section_IV_travaux(self):
        p = self.p
        env_parts = []
        if p.get("env_distance_centre"): env_parts.append(f"Distance centre : {p['env_distance_centre']}.")
        if p.get("env_acces"):           env_parts.append(f"Accès : {p['env_acces']}.")
        if p.get("env_transports"):      env_parts.append(f"Transports : {p['env_transports']}.")
        if p.get("env_commerces"):       env_parts.append(f"Commerces / services : {p['env_commerces']}.")
        env_text = " ".join(env_parts) or p.get("environnement_immediat", "")

        marche_text = ""
        if p.get("marche_tendance") or p.get("marche_tension"):
            marche_text = (
                f"Le marché immobilier local présente une tendance <b>{p.get('marche_tendance','Stable').lower()}</b> "
                f"avec une tension offre/demande <b>{p.get('marche_tension','équilibrée').lower()}</b> "
                f"au moment de l'expertise.")

        self.section("IV  —  TRAVAUX D'EXPERTISE")
        self.h2("1  —  Situation géographique et marché")
        self.para(p.get("situation_geographique",
            "Le bien immobilier se situe dans la commune de "
            f"{p.get('commune','')}, dans un secteur "
            f"{p.get('zone_urbanisme','urbanisé résidentiel')}. {env_text}"))
        if marche_text:
            self.para(marche_text)
        self.h2("2  —  Situation juridique")
        self.h3("Droit de propriété")
        self.para(p.get("origine_propriete",
            f"{p.get('demandeur_nom','')} est propriétaire du bien "
            "dont l'origine de propriété est précisée dans le titre transmis."))
        if any(p.get(k) for k in ("origine_nature", "origine_date_acquisition", "origine_prix_acquisition", "origine_notaire")):
            _prix = p.get("origine_prix_acquisition")
            self.kv([
                ("Nature de l'acquisition", p.get("origine_nature", "")),
                ("Date d'acquisition",      p.get("origine_date_acquisition", "") or "—"),
                ("Prix d'acquisition",      self._fmt(float(_prix)) if _prix else "—"),
                ("Notaire / étude",         p.get("origine_notaire", "") or "—"),
            ])
        self.h3("Données cadastrales et urbanisme")
        self.kv([
            ("Référence cadastrale", f"Section {p.get('cadastre_section','')} — N° {p.get('cadastre_num','')}"),
            ("Lieu-dit",             p.get("lieu_dit", "")),
            ("Contenance",           f"{p.get('terrain_m2','')} m²"),
            ("Zonage PLU",                  p.get("zonage_plu", "")),
            ("Constructibilité résiduelle", p.get("constructibilite") or "Non précisée"),
            ("Droit de préemption urbain",  p.get("dpu", "À vérifier")),
            ("Lot de copropriété",   p.get("lot_copro", "—") if p.get("type_bien") == "appartement" else "—"),
            ("Millièmes",            p.get("millesimes", "—") if p.get("type_bien") == "appartement" else "—"),
        ])
        self.h3("Servitudes, mitoyenneté et occupation")
        self.kv([
            ("Servitudes",             p.get("servitudes") or "Aucune servitude déclarée"),
            ("Mitoyenneté",            p.get("mitoyennete") or "Néant"),
            ("Situation d'occupation", p.get("situation_occupation", "") or "—"),
            ("Détail du bail",         p.get("bail_details") or "—"),
        ])
        self._section_IV_description()
        self._section_IV_risques()
        self._section_IV_swot()

    def _surfaces_kv(self):
        p = self.p
        rows = [("Surface de plancher (SDP)", f"{p.get('sdp','')} m²")]
        if p.get("surface_habitable"):
            rows.append(("Surface habitable", f"{p['surface_habitable']} m²"))
        if p.get("type_bien") == "appartement" and p.get("surface_carrez"):
            rows.append(("Surface loi Carrez", f"{p['surface_carrez']} m²"))
        self.kv(rows)

    def _section_IV_description(self):
        p = self.p
        t = p.get("type_bien", "maison")
        self.h2("3  —  Description des biens")

        if t == "maison":
            self.h3(f"3-1  La {p.get('type_bien_label','maison')}")
            self.para(p.get("description_construction",
                f"Construction datant de {p.get('annee_construction','')}. "
                f"État général : {p.get('etat_general','bon état')}. "
                f"Surface de plancher (SDP) : {p.get('sdp','')} m²."))
            self._surfaces_kv()
            self.h3("3-1-1  Disposition")
            for item in p.get("distribution", []):
                self.bullet(item)
            for extra in p.get("constructions_secondaires", []):
                self.h3(f"3-{extra.get('num','X')}  {extra.get('titre','')}")
                self.para(extra.get("description", ""))
                for item in extra.get("distribution", []):
                    self.bullet(item)
            self.h3("3-X  Le terrain d'assiette")
            self.para(p.get("description_terrain",
                f"Terrain d'une surface cadastrale de {p.get('terrain_m2','')} m². "
                f"{p.get('terrain_notes','')}"))

        elif t == "appartement":
            self.h3("3-1  L'appartement")
            self.para(p.get("description_construction",
                f"Appartement de type {p.get('type_appart','F3')} situé "
                f"{p.get('niveau_etage','au')} du bâtiment {p.get('batiment','')}. "
                f"État général : {p.get('etat_general','bon')}."))
            self._surfaces_kv()
            self.h3("3-1-1  Disposition")
            for item in p.get("distribution", []):
                self.bullet(item)
            if p.get("millesimes"):
                self.para(f"Et les {p['millesimes']} des parties communes générales de l'immeuble.")

        elif t == "terrain":
            self.h3("3-1  Le terrain")
            self.para(p.get("description_terrain",
                f"Terrain d'une surface de {p.get('terrain_m2','')} m². "
                f"{p.get('terrain_notes','')}"))

        carac = p.get("carac_physiques")
        if carac and t != "terrain":
            self.h3("3-X  État des composants — Constat lors de la visite")
            self._carac_physiques_table(carac)

        if t == "maison":
            postes_vet = p.get("vetuste_postes")
            if postes_vet and len(postes_vet) > 0:
                self.h3("3-X  Vétusté décomposée par poste (méthode pondérée)")
                self.legal(
                    "La vétusté est calculée poste par poste, pondérée par le "
                    "pourcentage de coût de chaque élément. La vétusté globale est la somme des contributions.")
                self._vetuste_table(postes_vet)

        self.legal(
            "Surface de Plancher (SDP) — art. R. 111-22 Code de l'urbanisme "
            "(décret 2011-2054 du 29/12/2011, en vigueur depuis le 1ᵉʳ mars 2012). "
            "La SHOB et la SHON sont des notions abrogées et ne peuvent plus être utilisées.")

    def _section_IV_risques(self):
        p = self.p
        self.h2("4  —  Analyse qualitative et risques")
        zone = p.get("zone_sismique", "V")
        risques = [
            ["Amiante",
             "Diagnostic requis" if not p.get("diag_amiante") else p.get("resultat_amiante", "Néant"),
             "Diagnostic à confier à un professionnel habilité." if not p.get("diag_amiante") else "Diagnostic fourni.",
             "CCH art. L. 271-4"],
            ["Termites",
             "Zone contaminée" if p.get("termites_zone", True) else "Hors zone",
             ("Diagnostic termites OBLIGATOIRE." if not p.get("diag_termites") else "Diagnostic fourni."),
             "Arrêté préfectoral Guadeloupe"],
            ["Sismicité", f"Zone {zone}",
             f"Zone {zone} — {'maximale (Guadeloupe)' if zone=='V' else 'forte'}. "
             "Conformité parasismique à vérifier pour les constructions antérieures à 2010.",
             "Décret 2010-1255"],
            ["Cyclones", "Zone exposée",
             "Territoire soumis au risque cyclonique — facteur intégré dans les prix de marché.",
             "PPRN applicable"],
            ["PPRN", "Applicable" if p.get("pprn", True) else "Non applicable",
             f"Plan de Prévention des Risques Naturels — commune de {p.get('commune','')}. "
             "Consulter le PPRN pour tout projet ou mise en conformité.",
             "Géorisques"],
        ]
        self.risque_table(risques)

    def _section_IV_swot(self):
        p = self.p
        forts = [l.strip() for l in (p.get("points_forts") or "").splitlines() if l.strip()]
        faibles = [l.strip() for l in (p.get("points_faibles") or "").splitlines() if l.strip()]
        if not forts and not faibles:
            return
        self.h2("5  —  Points forts et points faibles")
        self.para(
            "L'appréciation qualitative du bien fait ressortir les éléments suivants, qui ont été pris en "
            "compte dans la détermination de la valeur :")
        if forts:
            self.h3("Points forts")
            for f in forts:
                self.bullet(f)
        if faibles:
            self.h3("Points de vigilance")
            for f in faibles:
                self.bullet(f)

    # ─── SECTION V — ÉVALUATION ───────────────────────────────────────────────
    def _section_V_evaluation(self):
        t = self.p.get("type_bien", "maison")
        self.section("V  —  ÉVALUATION")
        self.para(
            "Compte tenu des éléments d'information mis à notre disposition et de notre visite "
            "des lieux, nous retenons les méthodes de calcul suivantes :")

        if t == "maison":
            self.para("— Méthode 1 : Par comparaison directe (DVF)", align='left', space_after=2)
            self.para("— Méthode 2 : Par sol et construction", align='left', space_after=2)
            if self.p.get("vente_anterieure"):
                self.para("— Méthode 3 : Par actualisation (indice ICC)", align='left', space_after=2)
            self._spacer(6)
            v1 = self._methode_comparaison()
            v2 = self._methode_sol_construction()
            v3 = self._methode_actualisation() if self.p.get("vente_anterieure") else None
            self._synthese_methodes(v1, v2, v3)

        elif t == "appartement":
            self.para("— Méthode unique : Par comparaison directe (DVF)", align='left', space_after=4)
            self.legal(
                "Pour un appartement en copropriété, la méthode par comparaison directe "
                "est la seule méthode appropriée. La méthode par sol et construction s'applique "
                "aux maisons individuelles. Aucun abattement de vétusté ne s'applique sur une "
                "valeur obtenue par comparaison : la vétusté est intégrée dans les prix de marché.")
            v1 = self._methode_comparaison()
            self._conclure_appart(v1)

        elif t == "terrain":
            self._methode_terrain()

    def _methode_comparaison(self):
        p = self.p
        refs = p.get("dvf_refs", [])
        pm2_median = p.get("pm2_median", 0)
        self.h2("5-1  Méthode par comparaison directe — Source DVF+")
        self.para(
            f"Source : Base DVF+ open-data (DGFiP — data.gouv.fr). "
            f"Transactions retenues : {p.get('type_bien_label','bien similaire')} "
            f"— secteur {p.get('commune','')}{' et communes comparables' if p.get('communes_comparables') else ''}. "
            f"Période : {p.get('periode_dvf','24 mois glissants')}.",
            size=9, color=MUTED, italic=True)
        if refs:
            self.dvf_table(refs, pm2_median)

        corrections = p.get("corrections_comparaison", [])
        if corrections:
            self.h3("Corrections appliquées :")
            corr_rows = []
            total_corr = 0
            for c in corrections:
                taux = c.get("taux", 0)
                total_corr += taux
                sign = "+" if taux >= 0 else ""
                corr_rows.append([c.get("critere", ""), f"{sign}{taux} %", c.get("justification", "")])
            pm2_corrige = round(pm2_median * (1 + total_corr / 100))
            corr_rows.append([
                "<b>Prix corrigé retenu</b>",
                f"<b>{pm2_corrige:,.0f}".replace(",", " ") + " €/m²</b>",
                f"Base {pm2_median:,.0f} €/m² × correction nette {'+' if total_corr>=0 else ''}{total_corr} %".replace(",", " "),
            ])
            self._styled_table(["Critère", "Taux", "Justification"], corr_rows,
                               [Cm(5.5), Cm(2.5), Cm(8.8)], total_row=True)
        else:
            pm2_corrige = pm2_median

        sdp = float(p.get("sdp", 0) or 0)
        val_comp = round(sdp * pm2_corrige / 1000) * 1000 if sdp else 0
        self.kv([
            ("SDP retenue",        f"{sdp} m²"),
            ("Prix corrigé /m²",   f"{pm2_corrige:,.0f}".replace(",", " ") + " €/m²"),
            ("<b>Valeur méthode 1</b>", f"<b>{self._fmt(val_comp)}</b>"),
        ])
        _med = f"{pm2_median:,.0f}".replace(",", " ")
        _corr = ""
        if corrections:
            _pc = f"{pm2_corrige:,.0f}".replace(",", " ")
            _corr = (f" Après application des corrections justifiées ci-dessus, le prix unitaire retenu "
                     f"s'établit à {_pc} €/m².")
        self.para(
            f"L'analyse des transactions comparables issues de la base DVF fait ressortir un prix de marché "
            f"médian de {_med} €/m².{_corr} Appliqué à la surface de plancher de {sdp:.0f} m² du bien "
            f"expertisé, ce prix conduit à une valeur par comparaison directe de {self._fmt(val_comp)}.")
        return val_comp

    def _methode_sol_construction(self):
        p = self.p
        self.h2("5-2  Méthode par sol et construction")
        terrain_m2   = float(p.get("terrain_m2", 0) or 0)
        prix_m2_brut = float(p.get("prix_terrain_brut", 120) or 120)
        decote_occu  = float(p.get("decote_occupation", 20) or 20) / 100
        vt_brut      = terrain_m2 * prix_m2_brut
        vt_nette     = round(vt_brut * (1 - decote_occu))
        self.h3("1) Valeur du terrain")
        self.kv([
            ("Surface cadastrale", f"{terrain_m2:.0f} m²"),
            ("Prix brut /m²",      f"{prix_m2_brut:.0f} €/m²"),
            ("Valeur brute",       self._fmt(vt_brut)),
            (f"Décote terrain occupé ({decote_occu*100:.0f} %)", f"− {self._fmt(vt_brut*decote_occu)}"),
            ("<b>Valeur terrain (VT)</b>", f"<b>{self._fmt(vt_nette)}</b>"),
        ])
        sdp     = float(p.get("sdp", 0) or 0)
        cout_m2 = float(p.get("cout_construction_m2", 1450) or 1450)
        vetuste = float(p.get("vetuste_taux", 20) or 20) / 100
        vc_neuf = sdp * cout_m2
        vc_nette = round(vc_neuf * (1 - vetuste))
        self.h3("2) Valeur de la construction")
        self.kv([
            ("Surface de plancher (SDP)",        f"{sdp:.0f} m²"),
            ("Coût de reconstruction au m² SDP", f"{cout_m2:.0f} €/m²  (indice BT01 Guadeloupe — 2026)"),
            ("Valeur à neuf",                    self._fmt(vc_neuf)),
            (f"Vétusté pondérée ({vetuste*100:.0f} %)", f"− {self._fmt(vc_neuf*vetuste)}"),
            ("Justification vétusté", p.get("justification_vetuste",
                f"Bâtiment de {datetime.date.today().year - int(p.get('annee_construction',2000) or 2000)} ans — vétusté pondérée par poste")),
            ("<b>Valeur construction (VC)</b>",  f"<b>{self._fmt(vc_nette)}</b>"),
        ])
        val_sc = round((vt_nette + vc_nette) / 1000) * 1000
        self.kv([
            ("VT (terrain)",      self._fmt(vt_nette)),
            ("VC (construction)", self._fmt(vc_nette)),
            ("<b>Valeur méthode 2</b>", f"<b>{self._fmt(val_sc)}</b>"),
        ])
        self.para(
            f"Cette méthode reconstitue la valeur du bien en additionnant la valeur du terrain "
            f"({self._fmt(vt_nette)}) et celle de la construction, après déduction de sa vétusté "
            f"({self._fmt(vc_nette)}), soit une valeur par sol et construction de {self._fmt(val_sc)}.")
        self.legal(
            "Note : le coût de 1 450 €/m² est calculé sur la Surface de Plancher (SDP) "
            "conformément à l'art. R. 111-22 du Code de l'urbanisme. "
            "La SHOB (Surface Hors Œuvre Brute) est une notion abrogée depuis mars 2012.")
        return val_sc

    def _methode_actualisation(self):
        p = self.p
        va = p.get("vente_anterieure", {})
        prix_base = float(va.get("prix", 0) or 0)
        indice_b  = float(va.get("indice_base", 1671) or 1671)
        indice_r  = float(va.get("indice_revision", 2146) or 2146)
        val_act   = round(prix_base * indice_r / indice_b / 1000) * 1000 if indice_b else 0
        self.h2("5-3  Méthode par actualisation — Indice ICC")
        self.para(
            f"Le présent bien a fait l'objet d'une vente le {va.get('date','')} "
            f"pour un montant de {self._fmt(prix_base)} net vendeur. "
            "L'actualisation s'effectue selon l'indice du coût de la construction (ICC — INSEE).")
        self.kv([
            ("Date de la vente antérieure", va.get("date", "")),
            ("Prix de référence",           self._fmt(prix_base)),
            ("Indice base (ICC)",           f"{indice_b:.0f}  —  {va.get('trimestre_base','')}"),
            ("Indice révision (ICC)",       f"{indice_r:.0f}  —  {va.get('trimestre_revision','')}"),
            ("Formule",                     f"{prix_base:.0f} × {indice_r:.0f} / {indice_b:.0f}"),
            ("<b>Valeur actualisée</b>",    f"<b>{self._fmt(val_act)}</b>"),
        ])
        return val_act

    def _methode_terrain(self):
        p = self.p
        terrain_m2  = float(p.get("terrain_m2", 0) or 0)
        prix_m2     = float(p.get("prix_terrain_brut", 50) or 50)
        val_terrain = round(terrain_m2 * prix_m2 / 1000) * 1000
        self.h2("5-1  Évaluation du terrain")
        self.kv([
            ("Surface cadastrale",   f"{terrain_m2:.0f} m²"),
            ("Prix unitaire retenu", f"{prix_m2:.0f} €/m²"),
            ("<b>Valeur terrain</b>", f"<b>{self._fmt(val_terrain)}</b>"),
        ])
        val_ret = val_terrain
        self.conclusion_box(round(val_ret*0.93/1000)*1000, val_ret, round(val_ret*1.05/1000)*1000)
        return val_ret

    def _synthese_methodes(self, v1, v2, v3=None):
        p = self.p
        methodes = []
        if v1: methodes.append(("M1 — Comparaison directe DVF", v1, p.get("poids_comparaison", 45)))
        if v2: methodes.append(("M2 — Sol & Construction", v2, p.get("poids_sol_construction", 35)))
        if v3: methodes.append(("M3 — Actualisation ICC", v3, p.get("poids_actualisation", 20)))
        total_poids = sum(m[2] for m in methodes) or 100
        val_pond = sum(m[1] * m[2] / total_poids for m in methodes)
        val_ret  = round(val_pond / 1000) * 1000

        decote_sci = p.get("decote_sci", 0)
        if decote_sci:
            val_avant = val_ret
            val_ret = round(val_ret * (1 - decote_sci / 100) / 1000) * 1000

        val_min = round(val_ret * 0.93 / 1000) * 1000
        val_max = round(val_ret * 1.05 / 1000) * 1000

        rows = [[m[0], self._fmt(m[1]), f"{m[2]} %", self._fmt(m[1]*m[2]/total_poids)] for m in methodes]
        rows.append(["<b>VALEUR RETENUE</b>", f"<b>{self._fmt(val_ret)}</b>", "100 %", f"<b>{self._fmt(val_ret)}</b>"])

        self.h2("5-X  Synthèse — Pondération des méthodes")
        self.synth_table(rows)
        self.para(
            "Les méthodes mises en œuvre sont convergentes. Après pondération tenant compte de la fiabilité "
            "respective de chacune au regard du bien et des données disponibles, nous arrêtons la valeur vénale "
            f"à {self._fmt(val_ret)}, dans une fourchette comprise entre {self._fmt(val_min)} et "
            f"{self._fmt(val_max)}.")
        if decote_sci:
            self.reserve(
                f"SCI — Décote de {decote_sci} % appliquée pour illiquidité des parts "
                f"({self._fmt(val_avant)} × {100-decote_sci} % = {self._fmt(val_ret)}). "
                "Fondée sur la doctrine BOFiP (IS - BIC Immobilier).",
                couleur=ORANGE, bg=ORANGEBG)
        self._spacer(8)
        self.conclusion_box(val_min, val_ret, val_max)
        return val_ret

    def _conclure_appart(self, val):
        val_min = round(val * 0.93 / 1000) * 1000
        val_max = round(val * 1.05 / 1000) * 1000
        self.conclusion_box(val_min, val, val_max)
        return val

    # ─── SECTION VI — CONCLUSIONS ─────────────────────────────────────────────
    def _build_conclusion_text(self):
        p = self.p
        sdp_part = f" — SDP {p['sdp']} m²" if p.get("sdp") else ""
        terrain_part = f" — terrain {p['terrain_m2']} m²" if p.get("terrain_m2") else ""
        return (
            f"Nous évaluons en définitive le bien immobilier situé "
            f"<b>{p.get('adresse_bien','')}, {p.get('commune','')} "
            f"({p.get('code_postal','')})</b> "
            f"({p.get('type_bien_label','')}{sdp_part}{terrain_part}), "
            f"en valeur vénale et en pleine propriété, "
            f"à la somme de <b>{p.get('valeur_retenue_texte','…')}.</b>")

    def _section_VI_conclusions(self):
        p = self.p
        self.section("VI  —  CONCLUSIONS")
        self.para(
            "Au terme de nos investigations et compte tenu de l'ensemble des éléments exposés dans le présent "
            "rapport — situation du bien, caractéristiques, état du marché local et résultats convergents des "
            "méthodes mises en œuvre — nous formulons la conclusion suivante.")
        self.para(self._build_conclusion_text())
        self.para("Les chiffres ci-dessus sont donnés hors fiscalité et hors frais d'acte.",
                  size=9, color=MUTED, italic=True)

        val_loc = float(p.get("valeur_locative_mensuelle") or 0)
        val_retenue_txt = p.get("valeur_retenue_texte", "")
        val_num = 0
        if val_retenue_txt:
            m = re.search(r"([\d\s  ]+)", val_retenue_txt.replace(",", ""))
            if m:
                try:
                    val_num = float(m.group(1).replace(" ", "").replace(" ", "").replace(" ", ""))
                except ValueError:
                    pass
        if val_loc > 0 and val_num > 0:
            taux_capi = val_loc * 12 / val_num * 100
            self.h3("Valeur locative de marché")
            self.kv([
                ("Loyer mensuel de marché estimé", f"{val_loc:,.0f} € / mois (charges non comprises)".replace(",", " ")),
                ("Revenu locatif annuel estimé",   f"{val_loc*12:,.0f} € / an".replace(",", " ")),
                ("Taux de capitalisation brut",    f"{taux_capi:.1f} %  (Marché Guadeloupe : 5–8 % brut typique)"),
            ])

        decote_liq = float(p.get("decote_liquidation") or 15)
        if val_num > 0 and decote_liq > 0:
            val_liq = round(val_num * (1 - decote_liq / 100) / 1000) * 1000
            self.h3("Valeur de liquidation rapide")
            self.para(
                f"Dans l'hypothèse d'une cession rapide (délai réduit — vente contrainte), "
                f"une décote de {decote_liq:.0f} % est usuellement retenue sur le marché guadeloupéen.")
            self.kv([
                ("Valeur vénale", self._fmt(val_num)),
                (f"Décote liquidation ({decote_liq:.0f} %)", f"− {self._fmt(val_num*decote_liq/100)}"),
                ("<b>Valeur de liquidation rapide</b>", f"<b>{self._fmt(val_liq)}</b>"),
            ])

        delai = p.get("delai_commercialisation")
        if delai:
            self.para(
                "Dans des conditions normales de marché, le délai de commercialisation du bien à la valeur "
                f"retenue est estimé à <b>{delai}</b>.")
        duree = p.get("duree_validite", "6 mois à compter de la date du rapport")
        self.legal(f"Durée de validité du présent rapport : {duree}.")

        if self._reserves:
            self.h3("Réserves émises :")
            for r, niveau in self._reserves:
                col = ALERT if niveau == "alerte" else NAVY
                self.reserve(r, couleur=col)

        self.para(
            "Arrêté le présent rapport d'expertise pour servir et valoir ce que de droit. "
            "Il a été établi en deux exemplaires originaux.")
        self._spacer(14)
        self.kv([
            ("Lieu et date",   f"Les Abymes (Guadeloupe), le {p.get('date_rapport','')}"),
            ("Expert",         f"{CABINET['expert']} — {CABINET['titre']}"),
            ("Cabinet",        f"{CABINET['nom']}  |  {CABINET['adresse']}"),
            ("Qualifications", f"{CABINET['rc']}  •  SIRET {CABINET['siret']}  •  RCP : {CABINET['rcp']}"),
        ])
        self._spacer(16)
        self.para("____________________________________", size=10, color=MUTED, align='right', space_after=1)
        self.para("Signature et cachet de l'expert", size=8, color=MUTED, align='right', space_after=6)
        self.legal(CABINET["refs_legales"])

    # ─── SECTION VII — ANNEXES ────────────────────────────────────────────────
    def _section_VII_annexes(self):
        p = self.p
        annexes = p.get("annexes", [
            "Titre de propriété",
            "Extrait du plan cadastral",
            "État des risques naturels et technologiques (ERP)",
            "Arrêté Préfectoral termites",
        ])
        self.section("VII  —  ANNEXES")
        self.h2("A — Reportage photographique")
        photos = p.get("photos")
        if photos and len(photos) > 0:
            self.para(
                f"Visite du {p.get('date_visite','')} — {len(photos)} photographie(s) jointe(s).",
                size=9, color=MUTED, italic=True)
            self._embed_photos(photos)
        else:
            self.para("Aucune photographie transmise dans le présent dossier.",
                      size=9, color=MUTED, italic=True)
        self.h2("B — Documents exploités")
        for i, a in enumerate(annexes, 1):
            self.para(f"{i}.  {a}", align='left', space_after=2)

    # ═════════════════════════════════════════════════════════════════════════
    # GÉNÉRATION
    # ═════════════════════════════════════════════════════════════════════════
    def generer(self, output_path: str) -> str:
        self._collecter_reserves()
        self._cover()
        self._section_I_serment()
        self._section_II_clause()
        self._section_III_mission()
        self._section_IV_travaux()
        self._section_V_evaluation()
        self._section_VI_conclusions()
        self._section_VII_annexes()
        self.doc.save(output_path)
        print(f"✓  Rapport Word généré : {output_path}")
        return output_path
